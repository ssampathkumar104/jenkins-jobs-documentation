import xml.etree.ElementTree as ET
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_jenkins_docx():
    xml_file = '4_ExecuteMultiplePlanesInParallel_job_config.xml'
    
    if not os.path.exists(xml_file):
        print(f"Error: '{xml_file}' not found in the current directory.")
        print("Please run your curl command first to download the file from AWS ALB.")
        return

    # Load and parse Jenkins XML configuration
    tree = ET.parse(xml_file)
    root = tree.getroot()

    # Initialize Word Document
    doc = Document()

    # --- Title Page / Header ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Jenkins Job Configuration Report")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(24)
    title_run.bold = True
    
    # Add a thin line spacing
    doc.add_paragraph()

    # --- Section 1: Description ---
    doc.add_heading("1. Job Overview & Description", level=1)
    description = root.find('description')
    desc_text = description.text if (description is not None and description.text) else 'No description provided.'
    doc.add_paragraph(desc_text)
    doc.add_paragraph() # Spacer

    # --- Section 2: Input Parameters ---
    doc.add_heading("2. Build Parameters & Input Variables", level=1)
    parameters = root.findall('.//hudson.model.ParametersDefinitionProperty/parameterDefinitions/')
    
    if parameters:
        for p in parameters:
            name_el = p.find('name')
            name = name_el.text if name_el is not None else "Unknown"
            
            default_val_el = p.find('defaultValue')
            val_text = default_val_el.text if default_val_el is not None else "None"
            
            # Format list item cleanly
            p_para = doc.add_paragraph(style='List Bullet')
            p_para.add_run("Parameter Name: ").bold = True
            p_para.add_run(f"{name}  |  ")
            p_para.add_run("Default Value: ").italic = True
            p_para.add_run(f"'{val_text}'").italic = True
    else:
        doc.add_paragraph("No user input parameters configured for this job.")
    doc.add_paragraph() # Spacer

    # --- Section 3: Build Scripts / Steps ---
    doc.add_heading("3. Execution Logic & Scripts", level=1)
    
    shell_steps = root.findall('.//hudson.tasks.Shell/command')
    pipeline_script = root.find('.//definition/script')

    # If it is a Freestyle job with Shell steps (Docker, Maven tasks, etc.)
    if shell_steps:
        for idx, step in enumerate(shell_steps, 1):
            doc.add_heading(f"Step {idx}: Shell Command Sequence", level=2)
            
            # Create a code block effect using indentation and Courier New font
            code_para = doc.add_paragraph()
            code_para.paragraph_format.left_indent = Inches(0.4)
            
            code_run = code_para.add_run(step.text.strip())
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(10)
            
    # If it is a Pipeline job (Jenkinsfile format)
    elif pipeline_script is not None and pipeline_script.text:
        doc.add_heading("Pipeline Groovy script (Jenkinsfile)", level=2)
        
        code_para = doc.add_paragraph()
        code_para.paragraph_format.left_indent = Inches(0.4)
        
        code_run = code_para.add_run(pipeline_script.text.strip())
        code_run.font.name = 'Courier New'
        code_run.font.size = Pt(10)
    else:
        doc.add_paragraph("No execution shell tasks or Groovy pipeline code scripts detected.")

    # Save the file
    output_filename = "Jenkins_Job_Documentation.docx"
    doc.save(output_filename)
    print(f"\nSuccess! Professional Word document generated: {output_filename}")

if __name__ == "__main__":
    create_jenkins_docx()